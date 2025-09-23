# modules/report.py
from __future__ import annotations

import math
from datetime import datetime
from pathlib import Path
from typing import Iterable, List, Dict, Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .utils import now_str, format_ts  # projedeki mevcut yardımcılar

# ------------------------------------------------------------
# Yardımcılar
# ------------------------------------------------------------
def _ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)

def _safe_write_text(path: Path, content: str) -> None:
    _ensure_parent(path)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content or "")

def _timestamp_suffix() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")

def _safe_docx_save(doc: Document, path: Path) -> Path:
    """Word dosyası açıksa PermissionError verir; timestamp'li yedekle kaydet."""
    _ensure_parent(path)
    try:
        doc.save(path)
        print(f"DOCX hazır → {path}")
        return path
    except PermissionError:
        alt = path.with_stem(f"{path.stem}_{_timestamp_suffix()}")
        doc.save(alt)
        print(f"DOCX farklı kaydedildi (dosya açıktı) → {alt}")
        return alt

def _dedupe_keep_order(lines: Iterable[str]) -> List[str]:
    seen, out = set(), []
    for x in lines or []:
        key = (x or "").strip()
        if key and key not in seen:
            seen.add(key)
            out.append(key)
    return out

def _srt_timestamp(total_seconds: float) -> str:
    """
    SubRip standardı: HH:MM:SS,mmm  (ör. 00:01:23,456)
    """
    if total_seconds is None:
        total_seconds = 0.0
    total_ms = int(round(float(total_seconds) * 1000))
    hours = total_ms // 3_600_000
    rest = total_ms % 3_600_000
    minutes = rest // 60_000
    rest = rest % 60_000
    seconds = rest // 1000
    millis = rest % 1000
    return f"{hours:02d}:{minutes:02d}:{seconds:02d},{millis:03d}"

def _window_label(window_summaries: Optional[List[Dict]]) -> str:
    """
    Pencere süresini (yaklaşık) etikete yansıt: 'Bölüm Özetleri (10 dk)' gibi.
    """
    if not window_summaries:
        return "Bölüm Özetleri"
    w = window_summaries[0]
    dur = max(1.0, float(w.get("end", 0) - w.get("start", 0)))
    minutes = int(round(dur / 60.0))
    return f"Bölüm Özetleri (~{minutes} dk)"

# ------------------------------------------------------------
# Düz metin çıktılar
# ------------------------------------------------------------
def save_transcript(text: str, filename: str = "transcript.txt"):
    path = Path(filename).absolute()
    _safe_write_text(path, text or "")
    print(f"Transcript kaydedildi → {path}")

def save_summary(summary: str, filename: str = "summary.txt"):
    path = Path(filename).absolute()
    _safe_write_text(path, summary or "")
    print(f"Özet kaydedildi → {path}")

def save_list(lines, filename: str):
    path = Path(filename).absolute()
    items = _dedupe_keep_order(lines or [])
    _safe_write_text(path, "\n".join(items))
    print(f"Liste kaydedildi → {path}")

# ------------------------------------------------------------
# SRT (subtitle) çıktısı
# ------------------------------------------------------------
def export_srt(segments: List[Dict], out: str = "meeting.srt"):
    path = Path(out).absolute()
    _ensure_parent(path)
    with open(path, "w", encoding="utf-8") as f:
        for i, s in enumerate(segments or [], 1):
            start = _srt_timestamp(float(s.get("start") or 0.0))
            end = _srt_timestamp(float(s.get("end") or 0.0))
            text = (s.get("text") or "").strip()
            if not text:
                continue
            f.write(f"{i}\n{start} --> {end}\n{text}\n\n")
    print(f"SRT kaydedildi → {path}")

# ------------------------------------------------------------
# Markdown notlar
# ------------------------------------------------------------
def export_notes_md(
    title: str,
    general_summary: str,
    window_summaries: Optional[List[Dict]],
    out: str = "notes.md",
):
    path = Path(out).absolute()
    _ensure_parent(path)
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"# {title}\n\n")
        f.write(f"**Tarih:** {now_str()}\n\n")
        f.write("## Genel Özet\n")
        f.write((general_summary or "-") + "\n\n")
        if window_summaries:
            f.write(f"## {_window_label(window_summaries)}\n")
            for w in window_summaries:
                st = _srt_timestamp(float(w.get("start") or 0.0))
                et = _srt_timestamp(float(w.get("end") or 0.0))
                f.write(f"- **{st}–{et}:** {w.get('summary') or '-'}\n")
    print(f"Markdown notlar → {path}")

# ------------------------------------------------------------
# DOCX (tutanak) üretimi
# ------------------------------------------------------------
def _apply_base_style(doc: Document):
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)

def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _add_bullet_list(doc: Document, items: Iterable[str]):
    for it in _dedupe_keep_order(items or []):
        doc.add_paragraph(it, style="List Bullet")

def _add_number_list(doc: Document, items: Iterable[str]):
    for it in _dedupe_keep_order(items or []):
        doc.add_paragraph(it, style="List Number")

def build_docx(
    meeting_title: str,
    summary: str,
    actions: Optional[List[str]],
    decisions: Optional[List[str]],
    segments: Optional[List[Dict]],
    out: str = "meeting_minutes.docx",
    window_summaries: Optional[List[Dict]] = None,
    max_segments_preview: int = 50,
) -> str:
    """
    Kurumsal formatta toplantı tutanağı üretir.
    - Yönetici özeti
    - Bölüm (pencere) özetleri
    - Aksiyon maddeleri
    - Alınan kararlar
    - Zaman çizelgesi (ilk N segment)
    """
    doc = Document()
    _apply_base_style(doc)

    # Başlık ve tarih
    _add_heading(doc, meeting_title, level=1)
    doc.add_paragraph(f"Tarih: {now_str()}")

    # Yönetici özeti
    _add_heading(doc, "Yönetici Özeti", level=2)
    doc.add_paragraph((summary or "-").strip())

    # Bölüm özetleri (varsa)
    if window_summaries:
        _add_heading(doc, _window_label(window_summaries), level=2)
        for w in window_summaries:
            st = _srt_timestamp(float(w.get("start") or 0.0))
            et = _srt_timestamp(float(w.get("end") or 0.0))
            doc.add_paragraph(f"[{st}–{et}] {w.get('summary') or '-'}")

    # Aksiyonlar
    _add_heading(doc, "Aksiyon Maddeleri", level=2)
    _add_bullet_list(doc, actions or ["-"])

    # Kararlar
    _add_heading(doc, "Alınan Kararlar", level=2)
    _add_bullet_list(doc, decisions or ["-"])

    # Zaman çizelgesi
    _add_heading(doc, f"Zaman Çizelgesi (ilk {max_segments_preview} segment)", level=2)
    for s in (segments or [])[:max_segments_preview]:
        st = _srt_timestamp(float(s.get("start") or 0.0))
        et = _srt_timestamp(float(s.get("end") or 0.0))
        line = (s.get("text") or "").strip()
        if line:
            doc.add_paragraph(f"[{st}–{et}] {line}")

    # Güvenli kaydetme
    final_path = _safe_docx_save(doc, Path(out).absolute())
    return str(final_path)
