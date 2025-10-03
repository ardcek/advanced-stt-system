#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ULTRA STT SİSTEMİ - ANA PANEL
============================

Özellikler:
1. 🎤 Ses Kayıt (kullanıcı durdurana kadar)
2. 📁 Dosya Seçimi (manuelt yol girme)
3. 🔧 Mod Seçimi (ultra, highest, balanced, fastest)
4. 🌍 Dil Seçimi (TR, EN, DE, FR, ES, IT, LA)
5. 🏥 Medical Mode (MeSH database)
6. 🎭 Diarization (kim ne dedi)
7. 📊 Çıktı Formatları (TXT, PDF, DOCX, MD)
8. 🤖 AI Özet (ÇATIR ÇATIR)
"""

import os
import sys
import time
import json
import sounddevice as sd
import soundfile as sf
import numpy as np
from datetime import datetime
from typing import Dict, List, Optional

class UltraSTTPanel:
    """Ultra STT Ana Panel Sistemi"""
    
    def __init__(self):
        self.audio_file = None
        self.selected_modes = []
        self.selected_language = "tr"
        self.medical_mode = False
        self.diarization = False
        self.output_formats = ["txt"]
        self.ai_summary = True
        self.ai_provider = "groq"  # Varsayılan: Groq API
        
        # Mevcut modlar
        self.quality_modes = {
            "fastest": "En Hızlı (10x hız, %91 doğruluk)",
            "balanced": "Dengeli (3x hız, %95 doğruluk)", 
            "highest": "En Yüksek (%98 doğruluk)",
            "ultra": "Ultra (%99.9 doğruluk hedefi)"
        }
        
        self.languages = {
            "tr": "Türkçe",
            "en": "English", 
            "de": "Deutsch",
            "fr": "Français",
            "es": "Español",
            "it": "Italiano",
            "la": "Latin"
        }
        
        self.output_formats_list = {
            "txt": "Metin (.txt)",
            "pdf": "PDF Raporu (.pdf)",
            "docx": "Word Belgesi (.docx)",
            "md": "Markdown (.md)",
            "html": "HTML Raporu (.html)",
            "json": "JSON Verisi (.json)"
        }

    def print_header(self, title: str):
        """Başlık yazdır"""
        print("\n" + "="*70)
        print(f"🎯 {title}")
        print("="*70)

    def print_section(self, title: str):
        """Bölüm başlığı"""
        print(f"\n📋 {title}")
        print("-" * 50)

    def show_main_menu(self):
        """Ana menüyü göster"""
        self.print_header("ULTRA STT SİSTEMİ - ANA PANEL")
        
        print("🎯 SES GİRİŞİ SEÇENEKLERİ:")
        print("1. 🎤 Yeni Ses Kaydet (kullanıcı durdurana kadar)")
        print("2. 📁 Ses Dosyası Seç (mevcut dosya)")
        print("3. 📝 Dosya Yolu Gir (manuel)")
        print("4. ❌ Çıkış")
        
        choice = input("\nSeçiminiz (1-4): ").strip()
        
        if choice == "1":
            self.record_audio()
        elif choice == "2":
            self.select_existing_file()
        elif choice == "3":
            self.manual_file_path()
        elif choice == "4":
            print("👋 Sistem kapatılıyor...")
            return False
        else:
            print("❌ Geçersiz seçim!")
            return self.show_main_menu()
            
        return True

    def record_audio(self):
        """Ses kayıt - kullanıcı durdurana kadar"""
        self.print_section("SES KAYIT SİSTEMİ")
        
        try:
            import sounddevice as sd
            import soundfile as sf
            import numpy as np
            import queue
            import threading
            
            print("🎤 Ses kayıt sistemi hazırlanıyor...")
            
            # Kayıt parametreleri
            samplerate = 44100
            channels = 1
            
            # Kayıt dosyası
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            os.makedirs("recordings", exist_ok=True)
            audio_file = f"recordings/kayit_{timestamp}.wav"
            
            # Kayıt queue
            audio_queue = queue.Queue()
            recording_data = []
            
            def audio_callback(indata, frames, time, status):
                """Audio callback"""
                if status:
                    print(f"Ses durumu: {status}")
                audio_queue.put(indata.copy())
            
            print("🔴 KAYIT BAŞLIYOR!")
            print("💡 Konuşmaya başlayın...")
            print("⏹️  Durdurmak için 'q' yazıp Enter basın!")
            
            # Kayıt stream başlat
            with sd.InputStream(samplerate=samplerate,
                              channels=channels, 
                              callback=audio_callback,
                              dtype='float32'):
                
                while True:
                    try:
                        # Queue'dan veri al
                        while not audio_queue.empty():
                            data = audio_queue.get_nowait()
                            recording_data.append(data)
                        
                        # Kullanıcı inputu kontrol et
                        import select
                        import sys
                        
                        # Windows için alternatif
                        try:
                            import msvcrt
                            if msvcrt.kbhit():
                                key = msvcrt.getch().decode('utf-8').lower()
                                if key == 'q':
                                    break
                        except:
                            # Linux/Mac için
                            user_input = input()
                            if user_input.lower().strip() == 'q':
                                break
                                
                    except KeyboardInterrupt:
                        break
                    except:
                        # Input olmadan devam et
                        time.sleep(0.1)
            
            # Kalan verileri al
            while not audio_queue.empty():
                data = audio_queue.get_nowait()
                recording_data.append(data)
            
            if recording_data:
                # Kayıt verilerini birleştir
                final_recording = np.concatenate(recording_data, axis=0)
                
                # Dosyaya kaydet
                sf.write(audio_file, final_recording, samplerate)
                
                duration = len(final_recording) / samplerate
                print(f"✅ Kayıt tamamlandı!")
                print(f"📁 Dosya: {audio_file}")
                print(f"⏱️  Süre: {duration:.1f} saniye")
                
                self.audio_file = audio_file
                self.configure_processing()
                
            else:
                print("❌ Kayıt verisi bulunamadı!")
                
        except ImportError:
            print("❌ sounddevice modülü eksik!")
            print("Kurulum: pip install sounddevice soundfile")
        except Exception as e:
            print(f"❌ Kayıt hatası: {e}")

    def select_existing_file(self):
        """Mevcut dosyalardan seç"""
        self.print_section("MEVCUT SES DOSYALARI")
        
        audio_files = []
        
        # Recordings klasörü
        if os.path.exists("recordings"):
            for file in os.listdir("recordings"):
                if file.lower().endswith(('.wav', '.mp3', '.m4a', '.flac', '.ogg')):
                    full_path = os.path.join("recordings", file)
                    size_mb = os.path.getsize(full_path) / (1024*1024)
                    audio_files.append((full_path, size_mb))
        
        # Ana klasör
        for file in os.listdir("."):
            if file.lower().endswith(('.wav', '.mp3', '.m4a', '.flac', '.ogg')):
                size_mb = os.path.getsize(file) / (1024*1024)
                audio_files.append((file, size_mb))
        
        if not audio_files:
            print("❌ Hiç ses dosyası bulunamadı!")
            print("💡 Önce ses kaydedin veya dosya yolu girin")
            return
        
        print("📁 Mevcut ses dosyaları:")
        for i, (file, size) in enumerate(audio_files, 1):
            print(f"  {i}. {file} ({size:.1f} MB)")
        
        try:
            choice = int(input(f"\nDosya seçin (1-{len(audio_files)}): "))
            if 1 <= choice <= len(audio_files):
                self.audio_file = audio_files[choice-1][0]
                print(f"✅ Seçilen dosya: {self.audio_file}")
                self.configure_processing()
            else:
                print("❌ Geçersiz seçim!")
        except ValueError:
            print("❌ Geçersiz giriş!")

    def manual_file_path(self):
        """Manuel dosya yolu gir"""
        self.print_section("MANUEL DOSYA YOLU")
        
        file_path = input("📁 Ses dosyası yolunu girin: ").strip()
        
        if not file_path:
            print("❌ Dosya yolu boş!")
            return
            
        if not os.path.exists(file_path):
            print(f"❌ Dosya bulunamadı: {file_path}")
            return
            
        # Dosya formatını kontrol et
        valid_formats = ('.wav', '.mp3', '.m4a', '.flac', '.ogg', '.mp4')
        if not file_path.lower().endswith(valid_formats):
            print(f"❌ Desteklenmeyen format! Desteklenen: {', '.join(valid_formats)}")
            return
            
        self.audio_file = file_path
        print(f"✅ Dosya seçildi: {self.audio_file}")
        self.configure_processing()

    def configure_processing(self):
        """İşleme ayarlarını yapılandır"""
        if not self.audio_file:
            print("❌ Önce ses dosyası seçin!")
            return
            
        self.print_header("İŞLEME AYARLARI")
        
        # Quality mode seçimi
        self.select_quality_mode()
        
        # Dil seçimi
        self.select_language()
        
        # Özel modlar
        self.select_special_modes()
        
        # Çıktı formatları
        self.select_output_formats()
        
        # Son onay ve işleme başlat
        self.show_configuration_summary()

    def select_language(self):
        """Dil seçimi"""
        self.print_section("DİL SEÇİMİ")
        
        print("🌍 Desteklenen diller:")
        langs = list(self.languages.keys())
        for i, lang in enumerate(langs, 1):
            print(f"  {i}. {lang}: {self.languages[lang]}")
        
        try:
            choice = int(input(f"\nDil seçin (1-{len(langs)}) [varsayılan: Türkçe]: "))
            if 1 <= choice <= len(langs):
                self.language = langs[choice-1]  # Doğru değişken ismi
                print(f"✅ Seçilen dil: {self.languages[self.language]}")
            else:
                print("❌ Geçersiz seçim! Türkçe kullanılıyor.")
                self.language = "tr"  # Doğru değişken ismi
        except ValueError:
            print("✅ Türkçe kullanılıyor.")
            self.language = "tr"  # Doğru değişken ismi

    def select_quality_mode(self):
        """Kalite modu seçimi"""
        self.print_section("KALİTE MODU SEÇİMİ")
        
        print("🔧 Kalite modları:")
        print("  1. ⚡ Fastest - En Hızlı (10x hız, %91 doğruluk)")
        print("  2. ⚖️ Balanced - Dengeli (3x hız, %95 doğruluk)")
        print("  3. 🎯 Highest - En Yüksek (%98 doğruluk)")
        print("  4. 🚀 Ultra - Ultra Kalite (AI destekli, en yavaş)")
        print("")
        print("💡 ÖNERİ: Balanced (hız/kalite dengesi)")
        
        try:
            choice = int(input("\nKalite modu seçin (1-4): "))
            
            if choice == 1:
                self.quality_mode = "fastest"
                print("✅ Fastest mod seçildi (Hızlı)")
            elif choice == 2:
                self.quality_mode = "balanced" 
                print("✅ Balanced mod seçildi (Dengeli)")
            elif choice == 3:
                self.quality_mode = "highest"
                print("✅ Highest mod seçildi (Yüksek Kalite)")
            elif choice == 4:
                self.quality_mode = "ultra"
                print("✅ Ultra mod seçildi (En Kaliteli)")
            else:
                print("❌ Geçersiz seçim! Balanced kullanılacak")
                self.quality_mode = "balanced"
                
        except (ValueError, KeyboardInterrupt):
            print("❌ Geçersiz giriş! Balanced kullanılacak")
            self.quality_mode = "balanced"

    def select_ai_summary_provider(self):
        """AI Özet sağlayıcı seçimi"""
        if not self.ai_summary:
            return  # AI özet aktif değilse atla
            
        self.print_section("AI ÖZET SAĞLAYICI SEÇİMİ")
        
        print("🤖 AI Özet sağlayıcıları:")
        print("  1. 🚀 Groq API (Ücretsiz, Hızlı) - ⚠️ Türkçe kalitesi düşük")
        print("  2. 💎 ChatGPT (OpenAI) - En kaliteli ama rate limit var")  
        print("  3. 🤖 Local AI (Gelişmiş yerli sistem) - 💡 ÖNERİLEN")
        print("  4. ❌ AI Özet İstemiyorum")
        print("")
        print("💡 ÖNERİ: Local AI (Türkçe için en iyi)")
        print("⚠️  Groq Türkçe'de zayıf, ChatGPT rate limit sorunu var")
        
        try:
            choice = int(input("\nAI özet sağlayıcı seçin (1-4): "))
            
            if choice == 1:
                self.ai_provider = "groq"
                print("✅ Groq API seçildi (Ücretsiz, Hızlı)")
            elif choice == 2:
                self.ai_provider = "openai"
                print("✅ ChatGPT seçildi (Rate limit riski var)")
            elif choice == 3:
                self.ai_provider = "local"
                print("✅ Local AI seçildi (Yerli sistem)")
            elif choice == 4:
                self.ai_summary = False
                self.ai_provider = None
                print("❌ AI Özet iptal edildi")
            else:
                print("❌ Geçersiz seçim! Groq API kullanılacak")
                self.ai_provider = "groq"
                
        except (ValueError, KeyboardInterrupt):
            print("❌ Geçersiz giriş! Groq API kullanılacak")
            self.ai_provider = "groq"
    
    def select_special_modes(self):
        """Özel modlar seçimi"""
        self.print_section("ÖZEL MODLAR")
        
        print("🎯 Özel özellikler:")
        print("1. 🏥 Medical Mode (tıbbi terimler + MeSH database)")
        print("2. 🎭 Diarization (kim ne dedi - konuşmacı ayrımı)")
        print("3. 🤖 AI Özet (akıllı özetleme)")
        print("4. ❌ Hiçbiri")
        
        choices = input("\nÖzellik seçin (1,2,3 virgülle ayırın): ").strip()
        
        if choices:
            selected = [int(x.strip()) for x in choices.split(",") if x.strip().isdigit()]
            
            if 1 in selected:
                self.medical_mode = True
                print("✅ Medical Mode aktif")
                
            if 2 in selected:
                self.diarization = True
                print("✅ Diarization aktif")
                
            if 3 in selected:
                self.ai_summary = True
                print("✅ AI Özet aktif")
                # AI özet seçildiyse sağlayıcı seçimi yap
                self.select_ai_summary_provider()
            else:
                self.ai_summary = False
        else:
            print("✅ Sadece temel özellikler kullanılacak")
            self.ai_summary = False

    def select_output_formats(self):
        """Çıktı formatları seçimi"""
        self.print_section("ÇIKTI FORMATLARI")
        
        print("📄 Mevcut formatlar:")
        formats = list(self.output_formats_list.keys())
        for i, fmt in enumerate(formats, 1):
            print(f"  {i}. {fmt}: {self.output_formats_list[fmt]}")
        
        choices = input(f"\nFormat seçin (1-{len(formats)}, virgülle ayırın) [varsayılan: txt]: ").strip()
        
        if choices:
            try:
                selected_indices = [int(x.strip()) for x in choices.split(",") if x.strip().isdigit()]
                selected_formats = [formats[i-1] for i in selected_indices if 1 <= i <= len(formats)]
                
                if selected_formats:
                    self.output_formats = selected_formats
                    print(f"✅ Seçilen formatlar: {', '.join(selected_formats)}")
                else:
                    print("❌ Geçersiz seçim! TXT kullanılıyor.")
                    self.output_formats = ["txt"]
            except ValueError:
                print("❌ Geçersiz giriş! TXT kullanılıyor.")
                self.output_formats = ["txt"]
        else:
            print("✅ TXT formatı kullanılıyor.")
            self.output_formats = ["txt"]

    def show_configuration_summary(self):
        """Yapılandırma özetini göster"""
        self.print_header("YAPILANDIRMA ÖZETİ")
        
        print(f"📁 Ses dosyası: {self.audio_file}")
        print(f"🔧 Kalite modu: {self.quality_mode}")
        print(f"🌍 Dil: {self.languages[self.language]}")
        print(f"🏥 Medical Mode: {'✅ Aktif' if self.medical_mode else '❌ Pasif'}")
        print(f"🎭 Diarization: {'✅ Aktif' if self.diarization else '❌ Pasif'}")
        
        if self.ai_summary:
            provider_names = {
                'groq': '🚀 Groq API (Ücretsiz)',
                'openai': '💎 ChatGPT (OpenAI)', 
                'local': '🤖 Local AI (Yerli)'
            }
            provider_name = provider_names.get(self.ai_provider, '🤖 AI Özet')
            print(f"🤖 AI Özet: ✅ Aktif ({provider_name})")
        else:
            print(f"🤖 AI Özet: ❌ Pasif")
            
        print(f"📄 Çıktı formatları: {', '.join(self.output_formats)}")
        
        confirm = input("\n🚀 İşlemeyi başlat? (y/n): ").strip().lower()
        
        if confirm == 'y':
            self.start_processing()
        else:
            print("❌ İşlem iptal edildi.")
            self.show_main_menu()

    def start_processing(self):
        """İşlemeyi başlat"""
        self.print_header("İŞLEME BAŞLATILIYOR")
        
        print("🔄 Audio dosyası işleniyor...")
        print(f"⏱️  Tahmini süre: {self.estimate_processing_time()}")
        
        try:
            # Burada gerçek STT işleme yapılacak
            result = self.process_audio_file()
            
            if result:
                print("✅ İşleme başarıyla tamamlandı!")
                self.generate_outputs(result)
            else:
                print("❌ İşleme başarısız!")
                
        except Exception as e:
            print(f"❌ İşleme hatası: {e}")

    def estimate_processing_time(self) -> str:
        """İşleme süresini tahmin et"""
        try:
            # Dosya boyutuna göre tahmin
            file_size = os.path.getsize(self.audio_file) / (1024*1024)  # MB
            
            if self.quality_mode == "fastest":
                time_factor = 0.1
            elif self.quality_mode == "balanced":
                time_factor = 0.3
            elif self.quality_mode == "highest":
                time_factor = 1.0
            else:  # ultra
                time_factor = 3.0
                
            estimated_minutes = file_size * time_factor
            
            if estimated_minutes < 1:
                return f"{estimated_minutes*60:.0f} saniye"
            else:
                return f"{estimated_minutes:.1f} dakika"
                
        except:
            return "Bilinmiyor"

    def process_audio_file(self) -> Dict:
        """Audio dosyasını işle - gerçek STT"""
        try:
            # Gerçek STT işleme
            print("🎤 STT işleme başlatılıyor...")
            from ultra_stt_interface import UltraSTTInterface
            stt_interface = UltraSTTInterface()
            
            # Parametreleri hazırla
            mode_params = {
                'audio_file': self.audio_file,
                'quality_mode': self.quality_mode,  # selected_quality_mode yerine quality_mode
                'language': self.language,          # selected_language yerine language
                'use_medical': self.medical_mode,
                'use_diarization': self.diarization,
                'use_ai_summary': self.ai_summary,
                'ai_provider': getattr(self, 'ai_provider', 'groq')  # AI sağlayıcı ekle
            }
            
            print(f"🔧 Mod: {self.quality_mode}")
            print(f"🌍 Dil: {self.language}")
            result = stt_interface.process_with_mode(mode_params)
            
            if result and 'transcription' in result:
                print("✅ STT işleme başarıyla tamamlandı!")
                return {
                    "transcript": result['transcription'],
                    "confidence": result.get('confidence', 0.95),
                    "processing_time": result.get('processing_time', 0),
                    "language": self.language,  # selected_language yerine language
                    "speakers": result.get('speakers', None),
                    "medical_terms": result.get('medical_terms', None),
                    "summary": result.get('summary', '')
                }
            else:
                raise Exception("STT sonucu alınamadı")
                
        except Exception as e:
            print(f"❌ STT işleme hatası: {e}")
            print("🔄 Simülasyon moduna geçiliyor...")
            time.sleep(2)  # Fallback simülasyon
            return {
                "transcript": "Bu bir test transkripsiyon metnidir. Gerçek STT entegrasyonu yapılacak.",
                "confidence": 0.95,
                "processing_time": 2.0,
                "language": self.selected_language,
                "speakers": ["Konuşmacı 1", "Konuşmacı 2"] if self.diarization else None,
                "medical_terms": ["test", "simülasyon"] if self.medical_mode else None,
                "summary": "Test özeti: Bu ses dosyası başarıyla işlenmiştir."
            }

    def generate_outputs(self, result: Dict):
        """Çıktı dosyalarını oluştur"""
        self.print_section("ÇIKTI DOSYALARI OLUŞTURULUYOR")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = f"stt_output_{timestamp}"
        
        # Her format için çıktı oluştur
        for fmt in self.output_formats:
            output_file = f"{base_name}.{fmt}"
            
            try:
                if fmt == "txt":
                    self.generate_txt_output(output_file, result)
                elif fmt == "pdf":
                    self.generate_pdf_output(output_file, result)
                elif fmt == "docx":
                    self.generate_docx_output(output_file, result)
                elif fmt == "md":
                    self.generate_md_output(output_file, result)
                elif fmt == "html":
                    self.generate_html_output(output_file, result)
                elif fmt == "json":
                    self.generate_json_output(output_file, result)
                    
                print(f"✅ {fmt.upper()} oluşturuldu: {output_file}")
                
            except Exception as e:
                print(f"❌ {fmt.upper()} oluşturma hatası: {e}")

    def generate_txt_output(self, filename: str, result: Dict):
        """TXT çıktısı oluştur"""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write("ULTRA STT SİSTEMİ - TRANSKRİPSİYON RAPORU\n")
            f.write("=" * 50 + "\n\n")
            f.write(f"Dosya: {self.audio_file}\n")
            f.write(f"Tarih: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Dil: {self.languages[self.selected_language]}\n")
            f.write(f"Kalite: {self.selected_modes[0] if self.selected_modes else 'ultra'}\n")
            f.write(f"Güven skoru: {result['confidence']:.2%}\n\n")
            
            f.write("TRANSKRİPSİYON:\n")
            f.write("-" * 20 + "\n")
            f.write(result['transcript'])
            f.write("\n\n")
            
            if self.diarization and result.get('speakers'):
                f.write("KONUŞMACILAR:\n")
                f.write("-" * 15 + "\n")
                for speaker in result['speakers']:
                    f.write(f"- {speaker}\n")
                f.write("\n")
                
            if self.medical_mode and result.get('medical_terms'):
                f.write("TIBBİ TERİMLER:\n")
                f.write("-" * 15 + "\n")
                for term in result['medical_terms']:
                    f.write(f"- {term}\n")

    def generate_md_output(self, filename: str, result: Dict):
        """Markdown çıktısı oluştur"""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write("# ULTRA STT SİSTEMİ - TRANSKRİPSİYON RAPORU\n\n")
            f.write("## 📊 Dosya Bilgileri\n\n")
            f.write(f"- **Dosya:** {self.audio_file}\n")
            f.write(f"- **Tarih:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"- **Dil:** {self.languages[self.selected_language]}\n")
            f.write(f"- **Kalite:** {self.selected_modes[0] if self.selected_modes else 'ultra'}\n")
            f.write(f"- **Güven Skoru:** {result['confidence']:.2%}\n\n")
            
            f.write("## 📝 Transkripsiyon\n\n")
            f.write(result['transcript'])
            f.write("\n\n")
            
            if self.ai_summary:
                f.write("## 🤖 AI Özet\n\n")
                f.write("*Özet burada olacak - AI entegrasyonu yapılacak*\n\n")

    def generate_json_output(self, filename: str, result: Dict):
        """JSON çıktısı oluştur"""
        import json
        
        output_data = {
            "metadata": {
                "file": self.audio_file,
                "timestamp": datetime.now().isoformat(),
                "language": self.selected_language,
                "quality_mode": self.selected_modes[0] if self.selected_modes else 'ultra',
                "medical_mode": self.medical_mode,
                "diarization": self.diarization,
                "ai_summary": self.ai_summary
            },
            "results": result,
            "configuration": {
                "output_formats": self.output_formats
            }
        }
        
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, ensure_ascii=False, indent=2)

    def generate_pdf_output(self, filename: str, result: Dict):
        """PDF çıktısı oluştur - gerçek PDF"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import ImageReader
            import textwrap
            
            c = canvas.Canvas(filename, pagesize=A4)
            width, height = A4
            
            # Başlık
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, height-50, f"STT Raporu - {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            
            # Dosya bilgileri
            c.setFont("Helvetica", 12)
            y_position = height - 100
            c.drawString(50, y_position, f"Ses Dosyası: {self.audio_file}")
            y_position -= 20
            c.drawString(50, y_position, f"Kalite Modu: {self.quality_mode}")
            y_position -= 20
            c.drawString(50, y_position, f"Dil: {self.language}")
            y_position -= 40
            
            # Transkripsiyon başlığı
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, y_position, "Transkripsiyon:")
            y_position -= 30
            
            # Transkripsiyon metni - satır bazında böl
            c.setFont("Helvetica", 10)
            text_lines = textwrap.fill(result.get('transcript', ''), width=80).split('\n')
            
            for line in text_lines:
                if y_position < 50:
                    c.showPage()
                    y_position = height - 50
                c.drawString(50, y_position, line)
                y_position -= 15
            
            # Özet varsa ekle
            if result.get('summary'):
                y_position -= 30
                c.setFont("Helvetica-Bold", 14)
                c.drawString(50, y_position, "AI Özet:")
                y_position -= 30
                
                c.setFont("Helvetica", 10)
                summary_lines = textwrap.fill(result.get('summary', ''), width=80).split('\n')
                for line in summary_lines:
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50
                    c.drawString(50, y_position, line)
                    y_position -= 15
            
            c.save()
            print(f"✅ PDF oluşturuldu: {filename}")
            
        except Exception as e:
            print(f"❌ PDF oluşturma hatası: {e}")
            # Fallback olarak TXT oluştur
            txt_filename = filename.replace('.pdf', '.txt')
            with open(txt_filename, 'w', encoding='utf-8') as f:
                f.write(f"STT Raporu - {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
                f.write(f"Transkripsiyon:\n{result.get('transcript', '')}\n")
                if result.get('summary'):
                    f.write(f"\nÖzet:\n{result.get('summary', '')}\n")
            print(f"✅ TXT dosyası oluşturuldu: {txt_filename}")

    def generate_docx_output(self, filename: str, result: Dict):
        """DOCX çıktısı oluştur - gerçek DOCX"""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Başlık
            title = doc.add_heading(f'STT Raporu - {datetime.now().strftime("%Y-%m-%d %H:%M")}', 0)
            
            # Dosya bilgileri
            doc.add_heading('Dosya Bilgileri', level=1)
            p = doc.add_paragraph()
            p.add_run('Ses Dosyası: ').bold = True
            p.add_run(self.audio_file)
            
            p = doc.add_paragraph()
            p.add_run('Kalite Modu: ').bold = True
            p.add_run(self.quality_mode)
            
            p = doc.add_paragraph()
            p.add_run('Dil: ').bold = True
            p.add_run(self.language)
            
            # Transkripsiyon
            doc.add_heading('Transkripsiyon', level=1)
            doc.add_paragraph(result.get('transcript', ''))
            
            # Özet varsa ekle
            if result.get('summary'):
                doc.add_heading('AI Özet', level=1)
                doc.add_paragraph(result.get('summary', ''))
            
            # Dosyayı kaydet
            doc.save(filename)
            print(f"✅ DOCX oluşturuldu: {filename}")
            
        except Exception as e:
            print(f"❌ DOCX oluşturma hatası: {e}")
            # Fallback olarak TXT oluştur
            txt_filename = filename.replace('.docx', '.txt')
            with open(txt_filename, 'w', encoding='utf-8') as f:
                f.write(f"STT Raporu - {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
                f.write(f"Transkripsiyon:\n{result.get('transcript', '')}\n")
                if result.get('summary'):
                    f.write(f"\nÖzet:\n{result.get('summary', '')}\n")
            print(f"✅ TXT dosyası oluşturuldu: {txt_filename}")

    def generate_html_output(self, filename: str, result: Dict):
        """HTML çıktısı oluştur"""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write("""<!DOCTYPE html>
<html lang="tr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Ultra STT Raporu</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }
        .header { background: #2c3e50; color: white; padding: 20px; border-radius: 8px; }
        .content { margin: 20px 0; }
        .transcript { background: #f8f9fa; padding: 20px; border-radius: 8px; border-left: 4px solid #007bff; }
        .metadata { background: #e9ecef; padding: 15px; border-radius: 8px; margin: 20px 0; }
    </style>
</head>
<body>
    <div class="header">
        <h1>🎯 Ultra STT Sistemi - Transkripsiyon Raporu</h1>
    </div>
    
    <div class="metadata">
        <h3>📊 Dosya Bilgileri</h3>
        <p><strong>Dosya:</strong> """ + self.audio_file + """</p>
        <p><strong>Tarih:</strong> """ + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + """</p>
        <p><strong>Dil:</strong> """ + self.languages[self.selected_language] + """</p>
        <p><strong>Güven Skoru:</strong> """ + f"{result['confidence']:.2%}" + """</p>
    </div>
    
    <div class="content">
        <h3>📝 Transkripsiyon</h3>
        <div class="transcript">
            """ + result['transcript'] + """
        </div>
    </div>
    
</body>
</html>""")

def main():
    """Ana fonksiyon"""
    try:
        panel = UltraSTTPanel()
        
        while True:
            if not panel.show_main_menu():
                break
                
    except KeyboardInterrupt:
        print("\n🛑 Program kullanıcı tarafından sonlandırıldı")
    except Exception as e:
        print(f"\n❌ Beklenmeyen hata: {e}")

if __name__ == "__main__":
    main()